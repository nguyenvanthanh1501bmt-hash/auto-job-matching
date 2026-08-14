from __future__ import annotations

import io
from dataclasses import dataclass

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen.canvas import Canvas


@dataclass(frozen=True, slots=True)
class GeneratedCvFixture:
    raw_cv_id: str
    filename: str
    content_type: str
    data: bytes
    expected_full_name: str
    expected_headline: str
    expected_normalized_job_title: str
    expected_skill: str


def generate_integration_fixtures(
) -> tuple[GeneratedCvFixture, ...]:
    return (
        GeneratedCvFixture(
            raw_cv_id=(
                "integration-software-engineer"
            ),
            filename=(
                "alex-morgan-software-engineer.pdf"
            ),
            content_type="application/pdf",
            data=_create_pdf(
                title=(
                    "Alex Morgan Software Engineer CV"
                ),
                lines=(
                    "ALEX MORGAN",
                    "Software Engineer",
                    (
                        "alex.morgan@example.com | "
                        "+1 415 555 0123"
                    ),
                    (
                        "Preferred Work Mode: "
                        "Remote, Hybrid"
                    ),
                    "",
                    "PROFESSIONAL SUMMARY",
                    (
                        "Software engineer experienced "
                        "in backend services and "
                        "business applications."
                    ),
                    "",
                    "SKILLS",
                    (
                        "Python, Java, Spring Boot, "
                        "Docker, PostgreSQL"
                    ),
                    "",
                    "WORK EXPERIENCE",
                    (
                        "Software Engineer | "
                        "Northstar Systems"
                    ),
                    "01/2021 - 06/2025",
                    (
                        "- Developed backend services "
                        "using Python and Java."
                    ),
                    (
                        "- Built business applications "
                        "with Spring Boot and PostgreSQL."
                    ),
                    (
                        "- Reduced batch processing "
                        "time by 30 percent."
                    ),
                    "",
                    "EDUCATION",
                    "Bachelor of Computer Science",
                    "City University",
                    "2016 - 2020",
                    "",
                    "LANGUAGES",
                    "English - Fluent",
                ),
            ),
            expected_full_name="ALEX MORGAN",
            expected_headline=(
                "Software Engineer"
            ),
            expected_normalized_job_title=(
                "SOFTWARE_ENGINEER"
            ),
            expected_skill="Python",
        ),
        GeneratedCvFixture(
            raw_cv_id=(
                "integration-senior-accountant"
            ),
            filename=(
                "anna-carter-senior-accountant.docx"
            ),
            content_type=(
                "application/vnd.openxmlformats-"
                "officedocument.wordprocessingml."
                "document"
            ),
            data=_create_docx(
                title=(
                    "Anna Carter Senior Accountant CV"
                ),
                lines=(
                    "ANNA CARTER",
                    "Senior Accountant",
                    (
                        "anna.carter@example.com | "
                        "+65 8123 4567"
                    ),
                    (
                        "Preferred Location: "
                        "Ho Chi Minh City"
                    ),
                    (
                        "Preferred Employment Type: "
                        "Full-time"
                    ),
                    "",
                    "PROFESSIONAL SUMMARY",
                    (
                        "Senior accountant experienced "
                        "in financial reporting, tax "
                        "accounting and auditing."
                    ),
                    "",
                    "SKILLS",
                    (
                        "Financial Reporting, "
                        "Tax Accounting, Auditing, "
                        "Budgeting, Microsoft Excel"
                    ),
                    "",
                    "WORK EXPERIENCE",
                    (
                        "Senior Accountant | "
                        "ABC Manufacturing"
                    ),
                    "01/2020 - 06/2024",
                    (
                        "- Prepared monthly and annual "
                        "financial statements."
                    ),
                    (
                        "- Managed tax declarations "
                        "and accounts payable."
                    ),
                    (
                        "- Reduced monthly closing "
                        "time by 25 percent."
                    ),
                    "",
                    "EDUCATION",
                    "Bachelor of Accounting",
                    "University of Economics",
                    "2015 - 2019",
                    "",
                    "CERTIFICATIONS",
                    "CPA",
                    (
                        "Issuer: Accounting "
                        "Professionals Association"
                    ),
                    "Issued: 2023-06",
                    "Expires: 2030-06",
                    "",
                    "LANGUAGES",
                    "English - Fluent",
                ),
            ),
            expected_full_name="ANNA CARTER",
            expected_headline=(
                "Senior Accountant"
            ),
            expected_normalized_job_title=(
                "SENIOR_ACCOUNTANT"
            ),
            expected_skill=(
                "Lập báo cáo tài chính"
            ),
        ),
        GeneratedCvFixture(
            raw_cv_id=(
                "integration-registered-nurse"
            ),
            filename=(
                "emily-johnson-registered-nurse.pdf"
            ),
            content_type="application/pdf",
            data=_create_pdf(
                title=(
                    "Emily Johnson Registered Nurse CV"
                ),
                lines=(
                    "EMILY JOHNSON",
                    "Registered Nurse",
                    (
                        "emily.johnson@example.com | "
                        "+44 7700 900123"
                    ),
                    (
                        "Preferred Employment Type: "
                        "Full-time"
                    ),
                    "",
                    "PROFESSIONAL SUMMARY",
                    (
                        "Registered nurse experienced "
                        "in patient care, clinical "
                        "assessment and infection control."
                    ),
                    "",
                    "SKILLS",
                    (
                        "Patient Care, Clinical Assessment, "
                        "Infection Control, Medical Records"
                    ),
                    "",
                    "WORK EXPERIENCE",
                    (
                        "Registered Nurse | "
                        "City General Hospital"
                    ),
                    "03/2020 - 06/2025",
                    (
                        "- Delivered patient care "
                        "across rotating clinical shifts."
                    ),
                    (
                        "- Performed clinical assessment "
                        "and maintained medical records."
                    ),
                    (
                        "- Improved compliance with "
                        "infection control procedures."
                    ),
                    "",
                    "EDUCATION",
                    "Bachelor of Nursing",
                    "University of Health Sciences",
                    "2015 - 2019",
                    "",
                    "LICENSES",
                    "Registered Nurse License",
                    (
                        "Issuing Authority: "
                        "National Nursing Council"
                    ),
                    "License Number: RN-2020-12345",
                    "Issued: 2020-01",
                    "Expires: 2030-01",
                    "",
                    "LANGUAGES",
                    "English - Native",
                ),
            ),
            expected_full_name="EMILY JOHNSON",
            expected_headline=(
                "Registered Nurse"
            ),
            expected_normalized_job_title=(
                "REGISTERED_NURSE"
            ),
            expected_skill="Chăm sóc bệnh nhân",
        ),
    )


def _create_pdf(
        *,
        title: str,
        lines: tuple[str, ...],
) -> bytes:
    output = io.BytesIO()

    canvas = Canvas(
        output,
        pagesize=A4,
        pageCompression=1,
    )

    canvas.setTitle(title)

    _, page_height = A4
    left_margin = 48
    top_margin = 48
    bottom_margin = 48
    line_height = 14

    y = page_height - top_margin

    canvas.setFont(
        "Helvetica",
        10,
    )

    for line in lines:
        if y <= bottom_margin:
            canvas.showPage()

            canvas.setFont(
                "Helvetica",
                10,
            )

            y = page_height - top_margin

        if line:
            canvas.drawString(
                left_margin,
                y,
                line,
            )

            y -= line_height
        else:
            y -= line_height // 2

    canvas.save()

    return output.getvalue()


def _create_docx(
        *,
        title: str,
        lines: tuple[str, ...],
) -> bytes:
    output = io.BytesIO()
    document = Document()

    document.core_properties.title = title

    section_headings = {
        "PROFESSIONAL SUMMARY",
        "CAREER OBJECTIVE",
        "SKILLS",
        "WORK EXPERIENCE",
        "EDUCATION",
        "CERTIFICATIONS",
        "LICENSES",
        "PROJECTS",
        "TRAINING",
        "LANGUAGES",
        "AWARDS",
        "PUBLICATIONS",
        "VOLUNTEERING",
        "ACTIVITIES",
        "INTERESTS",
    }

    for line in lines:
        if not line:
            document.add_paragraph()
            continue

        if line in section_headings:
            document.add_heading(
                line,
                level=1,
            )
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(
                style="List Bullet",
            )

            paragraph.add_run(
                line[2:]
            )

            continue

        document.add_paragraph(line)

    document.save(output)

    return output.getvalue()