from __future__ import annotations

import io
import subprocess
import zipfile

import fitz
import pytest
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen.canvas import Canvas

from app.config import Settings
from app.exceptions import (
    CvCorruptFileError,
    CvDocExtractionFailedError,
    CvExtractionTimeoutError,
    CvTextNotExtractableError,
    CvUnsupportedFormatError,
)
from app.extraction.doc_extractor import (
    OLE_COMPOUND_FILE_SIGNATURE,
    DocExtractor,
)
from app.extraction.docx_extractor import DocxExtractor
from app.extraction.extractor_factory import ExtractorFactory
from app.extraction.pdf_extractor import PdfExtractor


def _pdf_bytes(lines: list[str]) -> bytes:
    output = io.BytesIO()

    canvas = Canvas(
        output,
        pagesize=A4,
        pageCompression=1,
    )

    _, height = A4
    y = height - 48

    canvas.setFont(
        "Helvetica",
        10,
    )

    for line in lines:
        canvas.drawString(
            48,
            y,
            line,
        )
        y -= 14

    canvas.save()

    return output.getvalue()


def _blank_pdf_bytes() -> bytes:
    document = fitz.open()
    document.new_page()

    value = document.tobytes()

    document.close()

    return value


def _multi_column_pdf_bytes() -> bytes:
    document = fitz.open()

    page = document.new_page(
        width=600,
        height=800,
    )

    left_lines = [
        "Inventory management and inbound operations",
        "Warehouse safety and quality control",
        "Order fulfillment and stock reconciliation",
    ]

    right_lines = [
        "Team supervision across rotating shifts",
        "Forklift operation and equipment checks",
        "Reduced shipment damage by fifteen percent",
    ]

    for index, line in enumerate(left_lines):
        page.insert_text(
            (
                40,
                80 + index * 80,
            ),
            line,
            fontsize=10,
        )

    for index, line in enumerate(right_lines):
        page.insert_text(
            (
                340,
                80 + index * 80,
            ),
            line,
            fontsize=10,
        )

    value = document.tobytes()

    document.close()

    return value


def _docx_bytes() -> bytes:
    output = io.BytesIO()
    document = Document()

    document.add_paragraph(
        "NGUYỄN THỊ MINH AN"
    )

    document.add_paragraph(
        "Kế toán trưởng"
    )

    document.add_paragraph(
        "Báo cáo tài chính và kế toán thuế"
    )

    table = document.add_table(
        rows=2,
        cols=2,
    )

    table.cell(
        0,
        0,
    ).text = "Kỹ năng"

    table.cell(
        0,
        1,
    ).text = "Microsoft Excel"

    table.cell(
        1,
        0,
    ).text = "Ngôn ngữ"

    table.cell(
        1,
        1,
    ).text = "Tiếng Anh"

    section = document.sections[0]

    section.header.paragraphs[
        0
    ].text = "Candidate Profile"

    section.footer.paragraphs[
        0
    ].text = "AutoJob"

    document.save(output)

    return output.getvalue()


def _zip_bytes(
        entries: dict[str, bytes],
) -> bytes:
    output = io.BytesIO()

    with zipfile.ZipFile(
            output,
            mode="w",
            compression=zipfile.ZIP_DEFLATED,
    ) as archive:
        for name, value in entries.items():
            archive.writestr(
                name,
                value,
            )

    return output.getvalue()


def test_extracts_text_based_pdf(
        settings: Settings,
) -> None:
    extractor = PdfExtractor(settings)

    data = _pdf_bytes(
        [
            "ALEX MORGAN",
            "Senior Software Engineer",
            "Professional experience in Java, Python and Docker.",
            "Built reliable backend systems for business operations.",
        ]
    )

    result = extractor.extract(
        data,
        raw_cv_id="raw-pdf-1",
    )

    assert result.page_count == 1
    assert "ALEX MORGAN" in result.text
    assert "Senior Software Engineer" in result.text
    assert result.warnings == ()


def test_rejects_pdf_signature_mismatch(
        settings: Settings,
) -> None:
    extractor = PdfExtractor(settings)

    with pytest.raises(
            CvUnsupportedFormatError
    ) as captured:
        extractor.extract(
            b"not-a-pdf",
            raw_cv_id="raw-pdf-2",
        )

    assert (
            captured.value.code
            == "CV_UNSUPPORTED_FORMAT"
    )

    assert (
            captured.value.raw_cv_id
            == "raw-pdf-2"
    )


def test_rejects_corrupt_pdf(
        settings: Settings,
) -> None:
    extractor = PdfExtractor(settings)

    with pytest.raises(
            CvCorruptFileError
    ) as captured:
        extractor.extract(
            (
                b"%PDF-1.7\n"
                b"this is not a valid PDF body"
            ),
            raw_cv_id="raw-pdf-3",
        )

    assert (
            captured.value.code
            == "CV_CORRUPT_FILE"
    )


def test_rejects_pdf_without_extractable_text(
        settings: Settings,
) -> None:
    extractor = PdfExtractor(settings)

    with pytest.raises(
            CvTextNotExtractableError
    ) as captured:
        extractor.extract(
            _blank_pdf_bytes(),
            raw_cv_id="raw-pdf-4",
        )

    assert (
            captured.value.code
            == "CV_TEXT_NOT_EXTRACTABLE"
    )


def test_warns_for_multi_column_pdf(
        settings: Settings,
) -> None:
    extractor = PdfExtractor(settings)

    result = extractor.extract(
        _multi_column_pdf_bytes(),
        raw_cv_id="raw-pdf-5",
    )

    assert (
            "MULTI_COLUMN_LAYOUT_SUSPECTED"
            in result.warnings
    )

    assert (
            "TEXT_LAYOUT_MAY_BE_LOST"
            in result.warnings
    )


def test_extracts_docx_paragraphs_tables_header_footer(
        settings: Settings,
) -> None:
    extractor = DocxExtractor(settings)

    result = extractor.extract(
        _docx_bytes(),
        raw_cv_id="raw-docx-1",
    )

    assert (
            "NGUYỄN THỊ MINH AN"
            in result.text
    )

    assert (
            "Báo cáo tài chính"
            in result.text
    )

    assert (
            "Kỹ năng | Microsoft Excel"
            in result.text
    )

    assert (
            "Candidate Profile"
            in result.text
    )

    assert "AutoJob" in result.text
    assert result.page_count is None


def test_rejects_docx_signature_mismatch(
        settings: Settings,
) -> None:
    extractor = DocxExtractor(settings)

    with pytest.raises(
            CvUnsupportedFormatError
    ) as captured:
        extractor.extract(
            b"not-a-docx",
            raw_cv_id="raw-docx-2",
        )

    assert (
            captured.value.code
            == "CV_UNSUPPORTED_FORMAT"
    )


def test_rejects_corrupt_docx_archive(
        settings: Settings,
) -> None:
    extractor = DocxExtractor(settings)

    with pytest.raises(
            CvCorruptFileError
    ) as captured:
        extractor.extract(
            b"PK\x03\x04broken-archive",
            raw_cv_id="raw-docx-3",
        )

    assert (
            captured.value.code
            == "CV_CORRUPT_FILE"
    )


def test_rejects_office_archive_without_docx_document(
        settings: Settings,
) -> None:
    extractor = DocxExtractor(settings)

    data = _zip_bytes(
        {
            "[Content_Types].xml": (
                b"<Types />"
            ),
            "word/styles.xml": (
                b"<styles />"
            ),
        }
    )

    with pytest.raises(
            CvCorruptFileError
    ) as captured:
        extractor.extract(
            data,
            raw_cv_id="raw-docx-4",
        )

    assert (
            "does not contain a DOCX"
            in captured.value.message
    )


def test_rejects_docx_with_unsafe_archive_path(
        settings: Settings,
) -> None:
    extractor = DocxExtractor(settings)

    data = _zip_bytes(
        {
            "word/document.xml": (
                b"<document />"
            ),
            "../outside.txt": b"unsafe",
        }
    )

    with pytest.raises(
            CvCorruptFileError
    ) as captured:
        extractor.extract(
            data,
            raw_cv_id="raw-docx-5",
        )

    assert (
            "unsafe entry"
            in captured.value.message
    )


def test_rejects_docx_with_too_many_entries(
        settings: Settings,
) -> None:
    constrained = settings.model_copy(
        update={
            "max_docx_entries": 10,
        }
    )

    extractor = DocxExtractor(
        constrained
    )

    entries = {
        "word/document.xml": (
            b"<document />"
        ),
        **{
            f"word/item-{index}.xml": (
                b"<item />"
            )
            for index in range(10)
        },
    }

    with pytest.raises(
            CvCorruptFileError
    ) as captured:
        extractor.extract(
            _zip_bytes(entries),
            raw_cv_id="raw-docx-6",
        )

    assert (
            "too many entries"
            in captured.value.message
    )


def test_rejects_docx_exceeding_uncompressed_limit(
        settings: Settings,
) -> None:
    constrained = settings.model_copy(
        update={
            "max_docx_uncompressed_bytes": (
                1_048_576
            ),
        }
    )

    extractor = DocxExtractor(
        constrained
    )

    data = _zip_bytes(
        {
            "word/document.xml": (
                    b"A" * 1_048_577
            ),
        }
    )

    with pytest.raises(
            CvCorruptFileError
    ) as captured:
        extractor.extract(
            data,
            raw_cv_id="raw-docx-7",
        )

    assert (
            "uncompressed size limit"
            in captured.value.message
    )


def test_factory_selects_extractors_and_rejects_mismatch(
        settings: Settings,
) -> None:
    factory = ExtractorFactory(settings)

    pdf_extractor = factory.create(
        "candidate.pdf",
        "application/pdf; charset=binary",
        b"",
    )

    docx_extractor = factory.create(
        "candidate.docx",
        (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        b"",
    )

    doc_extractor = factory.create(
        "candidate.doc",
        "application/msword",
        b"",
    )

    assert (
            pdf_extractor.__class__.__name__
            == "PdfExtractor"
    )

    assert (
            docx_extractor.__class__.__name__
            == "DocxExtractor"
    )

    assert (
            doc_extractor.__class__.__name__
            == "DocExtractor"
    )

    with pytest.raises(
            CvUnsupportedFormatError
    ):
        factory.create(
            "candidate.txt",
            "text/plain",
            b"",
        )

    with pytest.raises(
            CvUnsupportedFormatError
    ):
        factory.create(
            "candidate.pdf",
            "application/msword",
            b"",
        )


def test_doc_reports_unavailable_antiword(
        settings: Settings,
        monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        (
            "app.extraction.doc_extractor."
            "shutil.which"
        ),
        lambda _: None,
    )

    extractor = DocExtractor(settings)

    with pytest.raises(
            CvDocExtractionFailedError
    ) as captured:
        extractor.extract(
            (
                    OLE_COMPOUND_FILE_SIGNATURE
                    + b"legacy-doc"
            ),
            raw_cv_id="raw-doc-1",
        )

    assert (
            captured.value.code
            == "CV_DOC_EXTRACTION_FAILED"
    )

    assert (
            "unavailable"
            in captured.value.message
    )


def test_doc_maps_antiword_timeout(
        settings: Settings,
        monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        (
            "app.extraction.doc_extractor."
            "shutil.which"
        ),
        lambda _: "antiword",
    )

    def raise_timeout(
            *args,
            **kwargs,
    ) -> None:
        del args
        del kwargs

        raise subprocess.TimeoutExpired(
            cmd="antiword",
            timeout=1,
        )

    monkeypatch.setattr(
        (
            "app.extraction.doc_extractor."
            "subprocess.run"
        ),
        raise_timeout,
    )

    extractor = DocExtractor(settings)

    with pytest.raises(
            CvExtractionTimeoutError
    ) as captured:
        extractor.extract(
            (
                    OLE_COMPOUND_FILE_SIGNATURE
                    + b"legacy-doc"
            ),
            raw_cv_id="raw-doc-2",
        )

    assert (
            captured.value.code
            == "CV_EXTRACTION_TIMEOUT"
    )


def test_doc_extracts_text_through_mocked_antiword(
        settings: Settings,
        monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        (
            "app.extraction.doc_extractor."
            "shutil.which"
        ),
        lambda _: "antiword",
    )

    monkeypatch.setattr(
        (
            "app.extraction.doc_extractor."
            "subprocess.run"
        ),
        lambda *args, **kwargs: (
            subprocess.CompletedProcess(
                args=args,
                returncode=0,
                stdout=(
                    "WAREHOUSE SUPERVISOR\n"
                    "Inventory management, "
                    "forklift operation and safety."
                ).encode("utf-8"),
                stderr=b"",
            )
        ),
    )

    extractor = DocExtractor(settings)

    result = extractor.extract(
        (
                OLE_COMPOUND_FILE_SIGNATURE
                + b"legacy-doc"
        ),
        raw_cv_id="raw-doc-3",
    )

    assert (
            "WAREHOUSE SUPERVISOR"
            in result.text
    )

    assert result.warnings == (
        "TEXT_LAYOUT_MAY_BE_LOST",
    )