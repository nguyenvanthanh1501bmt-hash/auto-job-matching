from __future__ import annotations

import io
import zipfile
from collections.abc import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.opc.exceptions import PackageNotFoundError
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from app.config import Settings
from app.exceptions import (
    CvCorruptFileError,
    CvTextNotExtractableError,
    CvUnsupportedFormatError,
)
from app.extraction.base import DocumentExtractor, ExtractionResult


# File DOCX thực chất là một ZIP archive chứa các XML document.
DOCX_REQUIRED_ENTRY = "word/document.xml"

# Magic bytes của ZIP archive.
ZIP_LOCAL_FILE_HEADER = b"PK\x03\x04"
ZIP_EMPTY_ARCHIVE = b"PK\x05\x06"
ZIP_SPANNED_ARCHIVE = b"PK\x07\x08"


class DocxExtractor(DocumentExtractor):
    def __init__(
            self,
            settings: Settings,
    ) -> None:
        self._settings = settings

    def extract(
            self,
            data: bytes,
            raw_cv_id: str | None = None,
    ) -> ExtractionResult:
        # Kiểm tra file có signature hợp lệ của ZIP/DOCX hay không.
        if not data.startswith(
                (
                        ZIP_LOCAL_FILE_HEADER,
                        ZIP_EMPTY_ARCHIVE,
                        ZIP_SPANNED_ARCHIVE,
                )
        ):
            raise CvUnsupportedFormatError(
                message=(
                    "The file extension or content type indicates DOCX, "
                    "but the file is not an Office Open XML archive"
                ),
                raw_cv_id=raw_cv_id,
            )

        # Kiểm tra cấu trúc và giới hạn an toàn của ZIP archive.
        self._validate_archive(
            data=data,
            raw_cv_id=raw_cv_id,
        )

        try:
            # Đọc DOCX trực tiếp từ bytes trong memory.
            document = Document(io.BytesIO(data))
        except (
                PackageNotFoundError,
                KeyError,
                ValueError,
                zipfile.BadZipFile,
                OSError,
        ) as exception:
            raise CvCorruptFileError(
                message="The DOCX document is corrupt or malformed",
                raw_cv_id=raw_cv_id,
            ) from exception

        text_parts: list[str] = []

        # Duyệt paragraph và table theo đúng thứ tự xuất hiện trong document.
        for block in self._iter_block_items(document):
            if isinstance(block, Paragraph):
                text = block.text.strip()

                if text:
                    text_parts.append(text)

                continue

            if isinstance(block, Table):
                text_parts.extend(
                    self._extract_table(block)
                )

        # Extract thêm nội dung từ header và footer của từng section.
        for section in document.sections:
            header_text = self._container_text(
                section.header
            )
            footer_text = self._container_text(
                section.footer
            )

            if header_text:
                text_parts.append(header_text)

            if footer_text:
                text_parts.append(footer_text)

        # Ghép toàn bộ text thành một chuỗi.
        text = "\n".join(text_parts).strip()

        # Chỉ tính chữ và số để xác định document có đủ nội dung hay không.
        meaningful = "".join(
            character
            for character in text
            if character.isalnum()
        )

        if len(meaningful) < self._settings.min_text_characters:
            raise CvTextNotExtractableError(
                raw_cv_id=raw_cv_id
            )

        return ExtractionResult(
            text=text,
            warnings=(),
            page_count=None,
        )

    # Kiểm tra cấu trúc và giới hạn an toàn của ZIP archive.
    def _validate_archive(
            self,
            data: bytes,
            raw_cv_id: str | None,
    ) -> None:
        try:
            with zipfile.ZipFile(
                    io.BytesIO(data)
            ) as archive:
                entries = archive.infolist()

                # Giới hạn số lượng entry để tránh archive quá lớn.
                if len(entries) > self._settings.max_docx_entries:
                    raise CvCorruptFileError(
                        message="The DOCX archive contains too many entries",
                        raw_cv_id=raw_cv_id,
                    )

                names = {
                    entry.filename
                    for entry in entries
                }

                # DOCX phải có document.xml chính.
                if DOCX_REQUIRED_ENTRY not in names:
                    raise CvCorruptFileError(
                        message=(
                            "The Office archive does not contain "
                            "a DOCX document"
                        ),
                        raw_cv_id=raw_cv_id,
                    )

                total_uncompressed = 0

                for entry in entries:
                    # Không cho phép encrypted archive.
                    if entry.flag_bits & 0x1:
                        raise CvCorruptFileError(
                            message=(
                                "Encrypted DOCX documents are not supported"
                            ),
                            raw_cv_id=raw_cv_id,
                        )

                    # Chặn path nguy hiểm như absolute path hoặc ../.
                    if self._is_unsafe_archive_path(
                            entry.filename
                    ):
                        raise CvCorruptFileError(
                            message=(
                                "The DOCX archive contains an unsafe entry"
                            ),
                            raw_cv_id=raw_cv_id,
                        )

                    # Theo dõi tổng kích thước sau khi giải nén.
                    total_uncompressed += entry.file_size

                    if (
                            total_uncompressed
                            > self._settings.max_docx_uncompressed_bytes
                    ):
                        raise CvCorruptFileError(
                            message=(
                                "The DOCX archive exceeds the configured "
                                "uncompressed size limit"
                            ),
                            raw_cv_id=raw_cv_id,
                        )

                    # Phát hiện compression ratio bất thường,
                    # giúp hạn chế ZIP bomb.
                    if (
                            entry.compress_size > 0
                            and entry.file_size > 10 * 1024 * 1024
                            and entry.file_size / entry.compress_size > 200
                    ):
                        raise CvCorruptFileError(
                            message=(
                                "The DOCX archive has a suspicious "
                                "compression ratio"
                            ),
                            raw_cv_id=raw_cv_id,
                        )

        except CvCorruptFileError:
            raise

        except zipfile.BadZipFile as exception:
            raise CvCorruptFileError(
                message="The DOCX document is corrupt or malformed",
                raw_cv_id=raw_cv_id,
            ) from exception

    @staticmethod
    def _is_unsafe_archive_path(
            filename: str,
    ) -> bool:
        # Chuẩn hóa path để kiểm tra cả Windows "\" và Unix "/".
        normalized = filename.replace(
            "\\",
            "/",
        )

        # Không cho phép absolute path.
        if normalized.startswith("/"):
            return True

        # Không cho phép path traversal bằng "..".
        return any(
            segment == ".."
            for segment in normalized.split("/")
        )

    @staticmethod
    def _iter_block_items(
            parent: DocumentObject | _Cell,
    ) -> Iterable[Paragraph | Table]:
        # Lấy XML body để duyệt paragraph/table theo thứ tự xuất hiện.
        parent_element = parent.element.body

        for child in parent_element.iterchildren():
            tag = child.tag

            if tag.endswith("}p"):
                yield Paragraph(
                    child,
                    parent,
                )

            elif tag.endswith("}tbl"):
                yield Table(
                    child,
                    parent,
                )

    @staticmethod
    def _extract_table(
            table: Table,
    ) -> list[str]:
        rows: list[str] = []

        # Duyệt từng row trong table.
        for row in table.rows:
            cells: list[str] = []
            seen_cells: set[int] = set()

            for cell in row.cells:
                # Một số merged cell có thể xuất hiện nhiều lần.
                cell_identity = id(cell._tc)

                if cell_identity in seen_cells:
                    continue

                seen_cells.add(cell_identity)

                # Ghép text của các paragraph trong cell.
                value = " ".join(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )

                if value:
                    cells.append(value)

            if cells:
                # Phân tách các cell trong cùng một row bằng "|".
                rows.append(
                    " | ".join(cells)
                )

        return rows

    @staticmethod
    def _container_text(container) -> str:
        # Extract text từ header/footer.
        values = [
            paragraph.text.strip()
            for paragraph in container.paragraphs
            if paragraph.text.strip()
        ]

        return "\n".join(values)